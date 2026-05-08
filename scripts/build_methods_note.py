"""
build_methods_note.py — Generate docs/synthesis-methods-note.docx

Synthesis Methods Note for africa-hiv-prep-atlas v0.1.0.
Format: Calibri 11pt, A4, 1.5 line spacing, 2.5cm margins.
Word limit: <=400 words (verified by tests/test_methods_note_word_count.py).
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

REPO = Path(__file__).resolve().parent.parent
OUTPUT = REPO / "docs" / "synthesis-methods-note.docx"

TITLE = (
    "African-cohort classification sensitivity in long-acting HIV PrEP "
    "meta-analyses: a systematic calibration study"
)

SECTIONS = {
    "Background": (
        "Long-acting HIV pre-exposure prophylaxis (LA-PrEP) trials have "
        "enrolled predominantly African participants, yet it is unknown how "
        "reliably published meta-analyses recognise and report this African "
        "cohort composition. We measured the sensitivity and specificity of "
        "MA-level African-cohort claims against a structured ground truth."
    ),
    "Methods": (
        "We identified 20 systematic reviews and meta-analyses of LA-PrEP "
        "efficacy published 2020-2026 via PubMed. For each of 54 (MA, trial) "
        "pairs we applied a three-layer algorithm (Layer-M: explicit count, "
        "table, and narrative) to classify whether the MA explicitly recognised "
        "the trial as African-cohort. Ground truth (D3) was defined as "
        ">=50% enrolment from African sites, derived from AACT structured "
        "country data for 7 LA-PrEP trials. Sensitivity and specificity were "
        "estimated with clustered bootstrap (n_clusters=20, 10,000 replicates). "
        "All code and fixtures are version-controlled at "
        "github.com/mahmood726-cyber/africa-hiv-prep-atlas (tag v0.1.0)."
    ),
    "Results": (
        "Against the D3 ground truth, MA African-cohort classification achieved "
        "35% sensitivity (95% CI 15-60%) and 71% specificity (95% CI 44-94%). "
        "The confusion matrix was: TP=13, FP=5, FN=24, TN=12. "
        "65% of actually-African LA-PrEP trials (24/37 MA-trial pairs) were "
        "cited but not classified as African across any of the three Layer-M "
        "layers. The pre-specified D2 sensitivity sweep (>=50% sites African) "
        "produced an identical confusion matrix to D3, confirming the result "
        "is invariant across reasonable site-share vs enrolment-share thresholds; "
        "D1 (>=1 African site) was non-discriminating because every fixture "
        "has at least one African site (TN=0). Cost-effectiveness MAs (2/20) "
        "contributed 5 false negatives through citation-by-number rather than "
        "narrative trial-naming. Pre-IRR working numbers; final values pending "
        "kappa >= 0.80 dual-rater audit (Task 20)."
    ),
    "Discussion": (
        "MAs systematically under-detect African cohort composition in "
        "LA-PrEP trials. The calibration gap is largest for economic and "
        "indirect-comparison MAs that cite trials by reference number. "
        "Standardised African-cohort reporting fields in SR protocols would "
        "reduce downstream classification error. Wide bootstrap confidence "
        "intervals reflect the small effective number of independent MAs; "
        "the blinded IRR audit will provide tighter estimates."
    ),
    "Data Availability": (
        "All data, code, and pre-registration materials are openly available "
        "at github.com/mahmood726-cyber/africa-hiv-prep-atlas (tag v0.1.0). "
        "atlas.csv is pinned byte-for-byte at the prereg-v0.1.0-amend-1 tag."
    ),
}


def _set_font(run, name: str = "Calibri", size_pt: int = 11):
    run.font.name = name
    run.font.size = Pt(size_pt)
    # Also set rPr/rFonts for complex script
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def _set_para_spacing(para, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)


def build(output_path: Path = OUTPUT) -> int:
    """Build the Methods Note. Returns word count."""
    doc = Document()

    # Page margins: 2.5cm all sides
    for section in doc.sections:
        section.page_height = Cm(29.7)   # A4
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_para = doc.add_heading(level=1)
    title_para.clear()
    run = title_para.add_run(TITLE)
    _set_font(run, "Calibri", 12)
    run.bold = True
    _set_para_spacing(title_para)

    # Body sections
    for heading, body in SECTIONS.items():
        h = doc.add_paragraph()
        run_h = h.add_run(heading)
        _set_font(run_h, "Calibri", 11)
        run_h.bold = True
        _set_para_spacing(h)

        p = doc.add_paragraph()
        run_b = p.add_run(body)
        _set_font(run_b, "Calibri", 11)
        _set_para_spacing(p)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Count words in body sections only (not title)
    all_text = " ".join(SECTIONS.values())
    word_count = len(all_text.split())
    return word_count


if __name__ == "__main__":
    wc = build()
    print(f"wrote {OUTPUT.relative_to(REPO)} ({wc} words)")
